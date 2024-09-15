# Loading Libraries 
import os
import glob
import docx
from tqdm import tqdm
import pandas as pd

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_openai import OpenAIEmbeddings
from langchain.storage import LocalFileStore
from langchain_community.vectorstores import FAISS
from langchain.embeddings import CacheBackedEmbeddings

# Setting Parameters
os.environ["OPENAI_API_KEY"] = "your-openai-api-key"

os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = "your_langchain_api_key_here"  

# Data  file names / directories
SourceFileDirectory = './source/'
FAISS_storage_directory = "./FAISS_storage/"
ChunksFileName = './intermediate_data/EU_REG_QA_Chunks.csv'

# Initialize Splitter 
CHUNK_SIZE = 2500
CHUNK_OVERLAP = 200
text_splitter = RecursiveCharacterTextSplitter(chunk_size = CHUNK_SIZE, chunk_overlap = CHUNK_OVERLAP, add_start_index=True)

# Initialize parameters of source documents / embedding model (this example uses OpenAI's embeddings)
EMBEDDING_MODELS_NAMES = ["text-embedding-3-small", "text-embedding-ada-002", "text-embedding-3-large"]
    
process_parameters = {
    'SourceFileName': ["EU_AI_ACT_2024_06_Q&A_DE.docx","EU_AI_ACT_2024_06_Q&A_EN.docx","EU_AI_ACT_2024_06_Q&A_GR.docx",
                       "EU_AI_ACT_2024_06_Q&A_HU.docx","EU_AI_ACT_2024_06_Q&A_IT.docx","EU_AI_ACT_2024_06_Q&A_RO.docx",
                       "EU_Late_Payment_2023_09_DE.docx","EU_Late_Payment_2023_09_EN.docx","EU_Late_Payment_2023_09_GR.docx",
                       "EU_Late_Payment_2023_09_HU.docx","EU_Late_Payment_2023_09_IT.docx","EU_Late_Payment_2023_09_RO.docx",
                       "EU_Single_Currency pack 2026 06_DE.docx","EU_Single_Currency pack 2026 06_EN.docx","EU_Single_Currency pack 2026 06_FR.docx"],
    'Regulation': ["EU_AI","EU_AI","EU_AI",
                   "EU_AI","EU_AI","EU_AI",
                   "EU_Late_Payment","EU_Late_Payment","EU_Late_Payment",
                   "EU_Late_Payment","EU_Late_Payment","EU_Late_Payment",
                   "EU_Single_Currency","EU_Single_Currency","EU_Single_Currency"],
    'Language': ["German","English","Greek","Hungarian","Italian","Romanian",
                 "German","English","Greek","Hungarian","Italian","Romanian",
                 "German","English","French"],
    }
# Initialize the tqdm library to add progress bars to loops and iterable objects
tqdm.pandas()

class regulation_chunks:
    """ Handling of regulations Q&A text chunks processing results """
    
    def __init__(self):
        self.chunk_num = 0
    
        self.regulation_chunks_data = {
            'Regulation': [], 
            'Language': [], 
            'Model': [],
            'FAISS_FileName': [], 
            'Question_No': [], 
            'Question': [],
            'Chunk_No': [],
            'Chunk': []
            }    
        
    def Save_to_csv(self):
        chunks_df = pd.DataFrame(self.regulation_chunks_data)
        try:
            chunks_df.to_csv(ChunksFileName, index=False)
            print(f"Data successfully saved to {ChunksFileName}")
        except IOError as e:
            print(f"Failed to write to {ChunksFileName}: {e}")
        except Exception as e:
            print(f"An unexpected error occurred: {e}")

def clean_FAISS_directory():              
    """Delete all files in the embeddings storage directory"""
    
    files = glob.glob(f"{FAISS_storage_directory}*")
    for file in files:
        os.remove(file)

def Load_Source_Data(SourceFile:str):
    """ Load Source Data """
    
    docs = ""
    try:
        doc = docx.Document(SourceFile)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)      
        docs = "\n".join(full_text)                 # Combine the lines into a single document
    except FileNotFoundError:
        print("Error: The file was not found.")
    except PermissionError:
        print("Error: You don't have permission to read the file.")
    except IsADirectoryError:
        print("Error: Expected a file, but found a directory.")
    except IOError as e:
        print(f"Error: An I/O error occurred: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    return docs

def Embed_and_Save_vectordata(docs,cached_embedder, FAISS_storage_file):
    """ Embedding through FAISS and storing the results """
    
    vector_store = FAISS.from_documents(docs, cached_embedder)         # Embedding through FAISS   
    vector_store.save_local(FAISS_storage_file, "index")               # Saving embedding results

def get_embeddings(text, row, embedding_model_name, FAISS_storage_file):
    """ Splits the regulation text into chunks and keeps information for relevant questions, regulation, model used and language written """
    
    regulations_chunks_handler.chunk_num = 0                                        # Initialize the counter of chunks
    chunks_contents = []                                                            # chunk contents of a document 
    
    questions = text.split('@@@')                                                  # Split the text in Questions based on the @@@ marker
    questions = [question.strip() for question in questions if question.strip()]   # Remove empty strings and strip whitespaces
    
    for question in questions:                                                     # Prepare the final list of articles  
        lines = question.split('\n')
    
        # Extract the question number and question text
        firstline  = lines[0].strip()                      # the first sentence of each question is of "Question x" form
        number = firstline.split()[1].strip()              # Assumes format "Question X" to get the number of the question
        if "Question" in firstline:                        #filtering  abnormalities in the annotation of questions
            question_text = lines[1].strip()               # Extract the question text from the third line 
            content = '\n'.join(lines[2:]).strip()         # Extract the answer text from the rest of the lines 
            chunks = text_splitter.split_text(content)     # Split the answer's content to chunks
            for chunk in chunks:                           # Update processing results
                """ for every chunk keeps necessary data in the form of a the regulation_chunks dictionary """        
                regulations_chunks_handler.regulation_chunks_data['Regulation'].append(row["Regulation"]), 
                regulations_chunks_handler.regulation_chunks_data['Language'].append(row["Language"]),  
                regulations_chunks_handler.regulation_chunks_data['Model'].append(embedding_model_name),
                regulations_chunks_handler.regulation_chunks_data['FAISS_FileName'].append(FAISS_storage_file)
                regulations_chunks_handler.regulation_chunks_data['Question_No'].append(number), 
                regulations_chunks_handler.regulation_chunks_data['Question'].append(question_text), 
                regulations_chunks_handler.regulation_chunks_data['Chunk_No'].append(regulations_chunks_handler.chunk_num),
                chunk_content = f"Answer_No: {number}\n Content: {chunk}"  #the chunks' text is refixed with Question Number and Title to provide reference     
                regulations_chunks_handler.regulation_chunks_data['Chunk'].append(chunk_content)  
                
                chunks_contents.append(chunk_content)
                regulations_chunks_handler.chunk_num += 1

    # return the documents to be embedded 
    docs = [Document(page_content=chunk) for chunk in chunks_contents] # Transformation from text to "documents" structure for the embedding process      
    return docs
    
def Process_Source_File(row):  
    """ Running the process for each Source file """

    #Load regulation Q&A source file
    SourceFileName = f"{SourceFileDirectory}{row["SourceFileName"]}"
    regulation_text = Load_Source_Data(SourceFileName)

    if regulation_text:                                                       # If Source File Load is successful     
        for embedding_model_name in EMBEDDING_MODELS_NAMES:                   # Running the process for each embedding model       

            # Initialize embedding model          
            embedding_model = OpenAIEmbeddings(model=embedding_model_name)
    
            # Initialize FAISS
            FAISS_storage_file = f"{FAISS_storage_directory}index_{row["Regulation"]}_{row["Language"]}_{embedding_model_name}"
            store = LocalFileStore(FAISS_storage_file)
            cached_embedder = CacheBackedEmbeddings.from_bytes_store(embedding_model, store, namespace=embedding_model.model)

            # Call the function to split the Source file text into chunks and get the chunks as documents 
            documents = get_embeddings(regulation_text, row, embedding_model_name, FAISS_storage_file)
            
            # Call the function to embed documents and store the vectordata in the FAISS directory
            Embed_and_Save_vectordata(documents, cached_embedder, FAISS_storage_file)

if __name__ == "__main__":

    """ Main Process """    
    clean_FAISS_directory()                                    # Initialize FAISS directory
    regulations_chunks_handler = regulation_chunks()           # Initialize an instance of the text chunks structure
    parameters_df = pd.DataFrame(process_parameters)           # Initialize the batch process information 
    parameters_df.progress_apply(Process_Source_File, axis=1)  # Run the batch process to read the Source files 
    regulations_chunks_handler.Save_to_csv()                   # Saving the instance of the text chunks structure
